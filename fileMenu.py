from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
import os.path
import shutil


class FileMenu:
    def __init__(self, root, mainMenu=None, mainFrame=None):
        self.commandArray = ['Новий Ctrl+N', 'Вiдкрити Ctrl+O', 'Зберегти Ctrl+S',
                             'Зберегти як...', 'Налаштування', 'Вийти']
        self.functionArray = [lambda e=None: self.__newFile(), lambda e=None: self.__openFile(),
                              lambda e=None: self.__saveFile(), lambda e=None: self.__saveAsFile(),
                              lambda e=None: self.__savePathOption(), lambda e=None: self.__exit()]

        self.commandDict = {}

        self.bindsCommandArray = ['<Control-n>', '<Control-o>', '<Control-s>']

        for index in range(len(self.commandArray)):
            self.commandDict[self.commandArray[index]] = self.functionArray[index]

        self.fileMenu = Menu(mainMenu, tearoff=0)

        for index, key in enumerate(self.commandDict):
            if index == 1 or index == 4 or index == 5:
                self.fileMenu.add_separator()
            self.fileMenu.add_command(label=self.commandArray[index],
                                      command=self.commandDict[self.commandArray[index]])

        mainMenu.add_cascade(label='Файл', menu=self.fileMenu)

        self.fileTypes = (("Documents Microsoft Word", "*.docx"), ("all types", "*.*"))
        self.textArea = mainFrame.textArea
        self.mainFrame = mainFrame
        self.root = root
        self.textLength = 0
        self.openFilePath = None
        self.mainName = 'Коректор'
        self.fileName = None

        for index in range(len(self.bindsCommandArray)):
            root.bind(self.bindsCommandArray[index], self.commandDict[self.commandArray[index]])

        self.root.protocol('WM_DELETE_WINDOW', self.__exit)

        with open('pathways', encoding='utf-8') as f:
            self.mainDirectory = f.readline()
            self.backUpDirectory = f.readline()

    def __openFile(self):
        if self.mainFrame.isTextChanged():
            answer = messagebox.askyesnocancel('Відкриття файла', 'Ви змінили текст. Бажаєте його зберегти?')
            if answer is None:
                pass
            elif answer:
                self.__saveFile()
                self.__open()
            elif not answer:
                self.__open()
        else:
            self.__open()

    def __saveFile(self):
        if self.fileName:
            self.__save(self.fileName.name)
        else:
            self.__saveAsFile()

    def __newFile(self):
        if self.mainFrame.isTextChanged():
            answer = messagebox.askyesnocancel('Новий файл', 'Ви змінили текст. Бажаєте його зберегти?')
            if answer is None:
                pass
            elif answer:
                self.__saveFile()
                self.mainFrame.clearTextArea()
                self.openFilePath = None
                self.root.title(self.mainName)
            elif not answer:
                self.mainFrame.clearTextArea()
                self.openFilePath = None
                self.root.title(self.mainName)
        else:
            self.mainFrame.clearTextArea()
            self.openFilePath = None
            self.root.title(self.mainName)

    def __exit(self):
        if self.mainFrame.isTextChanged():
            answer = messagebox.askyesnocancel('Вихід', 'Ви змінили текст. Бажаєте його зберегти?')
            if answer is None:
                pass
            elif answer:
                self.__saveFile()
                self.root.quit()
            elif not answer:
                self.root.quit()
        else:
            self.root.quit()

    def __saveAsFile(self):
        self.fileName = filedialog.asksaveasfile(filetypes=self.fileTypes, defaultextension='.docx')
        if self.fileName:
            self.__save(self.fileName.name)
            self.root.title(self.mainName + ' - ' + os.path.basename(self.fileName.name))

    def __save(self, pathToFile):
        self.mainFrame.resetTextLength()
        doc = Document()
        text = self.textArea.get('1.0', END).split('\n')
        for key, value in enumerate(text):
            doc.add_paragraph(value)
            if value.endswith('\n'):
                self.mainFrame.setTextLength(len(value) + 1)
            else:
                self.mainFrame.setTextLength(len(value))

        doc.save(str(pathToFile))
        messagebox.showinfo('Збереження файла', 'Файл збережен')

        self.mainFrame.resetTextSizeChanged()

    def __open(self):
        self.fileName = filedialog.askopenfile(filetypes=self.fileTypes)
        if not self.fileName.name.endswith('.docx'):
            messagebox.showwarning('', 'Формат файлу не підтримується')
            return

        if self.fileName:
            self.mainFrame.resetTextLength()
            self.mainFrame.clearTextArea()
            doc = Document(self.fileName.name)
            for key, paragraph in enumerate(doc.paragraphs):
                self.textArea.insert(INSERT, paragraph.text)
                if paragraph.text.endswith('\n'):
                    self.mainFrame.setTextLength(len(paragraph.text) + 1)
                else:
                    self.mainFrame.setTextLength(len(paragraph.text))

            self.root.title(self.mainName + ' - ' + os.path.basename(self.fileName.name))

    def __savePathOption(self):
        top = Toplevel(self.root)
        top.title('Путь сохранения файлов базы данных')
        top.resizable(False, False)

        insFrame = Frame(top)
        insFrame.pack(padx=7, pady=7)

        mainDirectory = Text(insFrame, width=45, height=2)
        mainDirectory.insert('1.0', self.mainDirectory)
        backUpDirectory = Text(insFrame, width=45, height=2)
        backUpDirectory.insert('1.0', self.backUpDirectory)

        Label(insFrame, text='Место основного хранения базы данных').grid(column=0, row=0)
        mainDirectory.grid(column=0, row=1)
        Label(insFrame, text='Место для резервного копирования').grid(column=0, row=2)
        backUpDirectory.grid(column=0, row=3)

        Button(insFrame, text='...', command=lambda x=mainDirectory, y=self.mainDirectory: self.__getMainDirectory(x)).grid(column=1, row=1)
        Button(insFrame, text='...', command=lambda x=backUpDirectory, y=self.backUpDirectory: self.__getBackUpDirectory(x)).grid(column=1, row=3)

        fr = Frame(insFrame)
        fr.grid(column=0, row=4)

        okButton = Button(fr, text='Ok', width=10, command=lambda x=top: self.__setToFile(x))
        cancelButton = Button(fr, text='Cancel', width=10, command=lambda: top.destroy())

        okButton.pack(side=LEFT)
        cancelButton.pack(side=LEFT)

    def __getMainDirectory(self, entry):
        text = filedialog.askdirectory()
        if text:
            self.mainDirectory = text + '/База даних Corrector'
            entry.delete('1.0', END)
            entry.insert(END, self.mainDirectory)

    def __getBackUpDirectory(self, entry):
        text = filedialog.askdirectory()
        if text:
            self.backUpDirectory = text + '/Резервна копія Corrector'
            entry.delete('1.0', END)
            entry.insert(END, self.backUpDirectory)

    def __setToFile(self, top):
        with open('pathways', 'w') as f:
            f.write(self.mainDirectory + '\n')
            f.write(self.backUpDirectory)

        top.destroy()
